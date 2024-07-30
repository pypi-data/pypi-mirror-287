import io
import json

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def _remove_row(table, row):
    """
    Remove line from table
    :param table: Table reference
    :param row: Row reference
    :return: None
    """
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)


def docx_update_datasheet(reg_map, data_sheet):
    """
    Update register map tables in existing data sheet
    :param reg_map: Json representation of modbus register map
    :param data_sheet: Data sheet document
    :return: None
    """
    mb = ModbusTable()
    mb.load(reg_map)
    mb.open_doc(data_sheet)
    mb.regs_table(reg_type='INPUT')
    mb.regs_table(reg_type='HOLD')
    mb.save(data_sheet)


class ModbusTable:
    """
    Class for printing and manipulating modbus register maps in MS Word document tables
    """
    def __init__(self):
        self.regs = None
        self.doc = None

    def load(self, filename):
        """
        Load modbus register map from json
        :param filename: Json file
        :return: None
        """
        with io.open(filename, 'r', encoding='utf-8-sig') as f:
            self.regs = json.load(f)

    def create_doc(self, title, font_size=10):
        """
        Create new MS Word document with given title and 'normal' font size
        :param title: Title of document (heading level 0)
        :param font_size: Font size in points
        :return: None
        """
        self.doc = Document()
        font = self.doc.styles['Normal'].font
        font.size = Pt(font_size)
        for section in self.doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        self.doc.add_heading(title, 0)

    def open_doc(self, filename):
        """
        Open already existing document
        :param filename: Name and path to document
        :return: None
        """
        self.doc = Document(filename)

    def save(self, filename):
        """
        Save manipulated document to file
        :param filename: Name of new file to save
        :return: None
        """
        self.doc.save(filename)

    def add_heading(self, heading="Heading", level=1):
        """
        Add heading to the end of currently manipulated document
        :param heading: Heading text
        :param level: Level 1 - 9 (0 is for title)
        :return: None
        """
        self.doc.add_heading(heading, level)

    def add_paragraph(self, text):
        """
        Add paragraph populated with given text the end of currently manipulated document
        :param text: Paragraph text
        :return: None
        """
        self.doc.add_paragraph(text)

    def find_table(self,  reg_type="INPUT"):
        """
        Find table within opened document for given type of modbus registers
        :param reg_type: Must be either 'INPUT' or 'HOLD'
        :return: Table reference on success
        :return: None on failure
        """
        idx = 0 if reg_type == "INPUT" else 1
        for table in self.doc.tables:
            header = table.rows[0].cells
            if len(header) == 4 and header[0].text == 'Address' and header[3].text == 'Description':
                if idx == 0:
                    for r in table.rows[1:]:
                        _remove_row(table, r)
                    return table
                else:
                    idx -= 1
        return None

    def regs_table(self, reg_type="INPUT"):
        """
        Fill register map of given type into table (either existing or newly created)
        :param reg_type: Must be either 'INPUT' or 'HOLD'
        :return: None
        """
        table = self.find_table(reg_type)

        if table is None:
            table = self.doc.add_table(rows=0, cols=0, style='Light Shading')

            table.add_column(360000 * 1)
            table.add_column(360000 * 2)
            table.add_column(360000 * 1)
            table.add_column(360000 * 20)
            header = table.add_row().cells

            header[0].text = 'Address'
            header[1].text = 'Name'
            header[2].text = 'Format'
            header[3].text = 'Description'

        for reg in self.regs:
            if reg["Type"] == reg_type:
                row = table.add_row().cells
                add = ""
                row[0].text = ', '.join(str(x) for x in reg["Address"])
                row[1].text = reg["Label"]
                row[2].text = reg["Format"]
                desc = row[3].paragraphs[0]
                desc.add_run(reg["Description"].rstrip().replace('\r', ''))

                sep = "\n"
                if reg["Access"] in ["RWF"]:
                    desc.add_run(f"{sep}Non-volatile, default: {reg['Value']}. ").bold = True
                    sep = ""
                if reg["Access"] in ["RWIF"]:
                    desc.add_run(f"{sep}One-time prog., default: {reg['Value']}. ").bold = True
                    sep = ""
                if reg["Access"] in ["RW"] and (reg["Min"] != 0 or reg["Max"] != 0):
                    desc.add_run(f"{sep}Default: {reg['Value']}. ").bold = True
                    sep = ""
                if reg["Min"] != 0 or reg["Max"] != 0:
                    desc.add_run(f"{sep}Minimum: {reg['Min']}. Maximum: {reg['Max']}. ").italic = True
                    sep = ""
                if reg["Unit"] != "":
                    desc.add_run(f"{sep}Unit: {reg['Unit']}. ")
                    sep = ""

                row[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

